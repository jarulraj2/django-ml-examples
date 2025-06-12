from django.shortcuts import render
from .forms import ResumeForm
from .utils import extract_resume_data
from django.http import HttpResponse
from django.template.loader import get_template
from django.template.loader import render_to_string

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

def resume_builder(request):
    data = None
    selected_template = "template1"
    if request.method == 'POST':
        form = ResumeForm(request.POST)
        if form.is_valid():
            user_text = form.cleaned_data['experience_text']
            selected_template = form.cleaned_data['template_choice']
            data = extract_resume_data(user_text)
    else:
        form = ResumeForm()

    return render(request, 'generator/dashboard.html', {
        'form': form,
        'data': data,
        'template': selected_template,
    })

def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)  # e.g., "BDD7EE" is light blue
    tcPr.append(shd)

def generate_docx(request):
    if request.method == 'POST':
        user_text = request.POST.get("experience_text", "")
        selected_template = request.POST.get("template_choice", "template1")
        data = extract_resume_data(user_text)

        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # 🌟 Template 1
        if selected_template == "template1":
            # 🧑 Name
            doc.add_heading(data.get("name", "Your Name"), level=0)

            # 📧 Contact Info
            email = data.get("email", "you@example.com")
            phone = data.get("phone", "000-000-0000")
            doc.add_paragraph(f"Email: {email} | Phone: {phone}")

            # 🎯 Job Role
            doc.add_heading("Job Role", level=1)
            doc.add_paragraph(data.get("job_role", "Desired Role"))

            # 🧠 Skills
            skills = data.get("skills", [])
            doc.add_heading("Skills", level=2)
            doc.add_paragraph(", ".join(skills) if skills else "Not provided")

            # 💼 Experience
            experiences = data.get("experience", [])
            doc.add_heading("Experience", level=2)
            if experiences:
                for exp in experiences:
                    position = exp.get("position", "Position")
                    company = exp.get("company", "Company")
                    from_date = exp.get("from", "From")
                    to_date = exp.get("to", "To")
                    doc.add_paragraph(
                        f"{position} at {company} ({from_date} - {to_date})",
                        style='List Bullet'
                    )
            else:
                doc.add_paragraph("No experience provided.")

            # 🎓 Education
            education = data.get("education", [])
            doc.add_heading("Education", level=2)
            if education:
                for edu in education:
                    doc.add_paragraph(edu, style="List Bullet")
            else:
                doc.add_paragraph("No education info provided.")

            # 📜 Certificates
            certificates = data.get("certificates", [])
            doc.add_heading("Certificates", level=2)
            if certificates:
                for cert in certificates:
                    doc.add_paragraph(cert, style="List Bullet")
            else:
                doc.add_paragraph("No certifications provided.")

            # 💡 Projects
            projects = data.get("projects", [])
            doc.add_heading("Projects", level=2)
            if projects:
                for proj in projects:
                    doc.add_paragraph(f"• {proj}", style="List Bullet")
            else:
                doc.add_paragraph("No projects listed.")

        # 🌟 Template 2
        elif selected_template == "template2":
            doc.add_paragraph(data.get("name", "Your Name"), style='Title')
            doc.add_paragraph(f"Contact: {data.get('email')} / {data.get('phone')}")
            doc.add_paragraph("Role: " + data.get("job_role", "N/A"))

            doc.add_paragraph("Skills:", style="Heading 2")
            for skill in data.get("skills", []):
                doc.add_paragraph(f"- {skill}", style="List Bullet")

            doc.add_paragraph("Education:", style="Heading 2")
            for edu in data.get("education", []):
                doc.add_paragraph(f"- {edu}", style="List Bullet")

            doc.add_paragraph("Certificates:", style="Heading 2")
            for cert in data.get("certificates", []):
                doc.add_paragraph(f"- {cert}", style="List Bullet")

            doc.add_paragraph("Experience:", style="Heading 2")
            for exp in data.get("experience", []):
                doc.add_paragraph(f"{exp.get('position')} at {exp.get('company')} ({exp.get('from')} - {exp.get('to')})", style="List Bullet")

            doc.add_paragraph("Projects:", style="Heading 2")
            for proj in data.get("projects", []):
                doc.add_paragraph(f"- {proj}", style="List Bullet")


        # 🌟 Template 3
        elif selected_template == "template3":
            doc.add_heading(data.get("name", "Your Name"), level=1)
            doc.add_paragraph(f"{data.get('job_role', '')} | {data.get('email', '')} | {data.get('phone', '')}")

            doc.add_paragraph("Skills:", style="Heading 2")
            for skill in data.get("skills", []):
                doc.add_paragraph(f"• {skill}", style="List Bullet")

            doc.add_paragraph("Education:", style="Heading 2")
            for edu in data.get("education", []):
                doc.add_paragraph(f"• {edu}", style="List Bullet")

            doc.add_paragraph("Certificates:", style="Heading 2")
            for cert in data.get("certificates", []):
                doc.add_paragraph(f"• {cert}", style="List Bullet")

            doc.add_paragraph("Experience:", style="Heading 2")
            for exp in data.get("experience", []):
                doc.add_paragraph(f"{exp.get('position')} at {exp.get('company')} ({exp.get('from')} - {exp.get('to')})", style="List Bullet")

            doc.add_paragraph("Projects:", style="Heading 2")
            for proj in data.get("projects", []):
                doc.add_paragraph(f"• {proj}", style="List Bullet")


        # 🌟 Template 4
        elif selected_template == "template4":
            doc.add_paragraph(f"{data.get('name')} - {data.get('job_role')}", style='Title')
            doc.add_paragraph(f"Email: {data.get('email')}")
            doc.add_paragraph(f"Phone: {data.get('phone')}")

            doc.add_paragraph("Skills:", style="Heading 2")
            for skill in data.get("skills", []):
                doc.add_paragraph(skill, style="List Bullet")

            doc.add_paragraph("Education:", style="Heading 2")
            for edu in data.get("education", []):
                doc.add_paragraph(edu, style="List Bullet")

            doc.add_paragraph("Certificates:", style="Heading 2")
            for cert in data.get("certificates", []):
                doc.add_paragraph(cert, style="List Bullet")

            doc.add_paragraph("Experience:", style="Heading 2")
            for exp in data.get("experience", []):
                doc.add_paragraph(f"{exp.get('position')} at {exp.get('company')} ({exp.get('from')} - {exp.get('to')})", style="List Bullet")

            doc.add_paragraph("Projects:", style="Heading 2")
            for proj in data.get("projects", []):
                doc.add_paragraph(proj, style="List Bullet")

                            

        # 🌟 Template 5 – 2-column layout with colored left panel
        elif selected_template == "template5":
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            left_cell, right_cell = table.rows[0].cells

            # 🎨 Left Panel – dark background + white text
            set_cell_background(left_cell, "2C3E50")
            white = RGBColor(255, 255, 255)

            # 🧑 Name
            left_p = left_cell.paragraphs[0]
            run = left_p.add_run(data["name"])
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = white

            # 📧 Email & Phone
            for label in [f"Email: {data['email']}", f"Phone: {data['phone']}"]:
                para = left_cell.add_paragraph()
                run = para.add_run(label)
                run.font.color.rgb = white

            # 🧠 Skills
            para = left_cell.add_paragraph()
            run = para.add_run("Skills")
            run.bold = True
            run.font.color.rgb = white

            for skill in data["skills"]:
                para = left_cell.add_paragraph(style="List Bullet")
                run = para.add_run(skill)
                run.font.color.rgb = white

            # 🎓 Education
            para = left_cell.add_paragraph()
            run = para.add_run("Education")
            run.bold = True
            run.font.color.rgb = white

            for edu in data["education"]:
                para = left_cell.add_paragraph(style="List Bullet")
                run = para.add_run(edu)
                run.font.color.rgb = white

            # 📜 Certificates
            para = left_cell.add_paragraph()
            run = para.add_run("Certificates")
            run.bold = True
            run.font.color.rgb = white

            for cert in data["certificates"]:
                para = left_cell.add_paragraph(style="List Bullet")
                run = para.add_run(cert)
                run.font.color.rgb = white

            # 📝 Right Panel
            para = right_cell.paragraphs[0]
            run = para.add_run("Job Role")
            run.bold = True

            right_cell.add_paragraph(data["job_role"])

            para = right_cell.add_paragraph()
            run = para.add_run("Experience")
            run.bold = True

            for exp in data["experience"]:
                para = right_cell.add_paragraph(
                    f"{exp['position']} at {exp['company']} ({exp['from']} - {exp['to']})",
                    style="List Bullet"
                )

            para = right_cell.add_paragraph()
            run = para.add_run("Projects")
            run.bold = True

            for proj in data["projects"]:
                para = right_cell.add_paragraph(f"• {proj}", style="List Bullet")



        # 📥 Return Word doc as downloadable file
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        filename = f'{data["name"].replace(" ", "_")}_resume.docx'
        response['Content-Disposition'] = f'attachment; filename={filename}'
        doc.save(response)
        return response